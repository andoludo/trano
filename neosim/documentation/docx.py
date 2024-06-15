from pathlib import Path
from typing import Any, Dict, List, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt
from docx.table import _Cell

from neosim.documentation.documentation import ModelDocumentation, get_description

COLUMN_SIZE_WITH_DESCRIPTION = 3

TABLE_COUNT = 1


def _round(value: float | str | None) -> str:
    try:
        return str(round(value, 4))  # type: ignore
    except TypeError:
        return str(value)


def _set_color(cell: _Cell, key: str) -> None:
    if key == "type":
        cell._tc.get_or_add_tcPr().append(
            parse_xml(rf'<w:shd {nsdecls("w")} w:fill="#D3D3D3"/>')
        )
    if key == "name":
        cell._tc.get_or_add_tcPr().append(
            parse_xml(rf'<w:shd {nsdecls("w")} w:fill="#FF7F7F"/>')
        )
    if key == "thickness":
        cell._tc.get_or_add_tcPr().append(
            parse_xml(rf'<w:shd {nsdecls("w")} w:fill="#FFFFE0"/>')
        )


def create_table(doc: Document, data: Dict[str, Any]) -> None:  # type: ignore
    size = 2
    description = get_description()
    if any(key in description for key in data):
        size = COLUMN_SIZE_WITH_DESCRIPTION
    table = doc.add_table(cols=size, rows=len(data))  # type: ignore
    for i, (key, value) in enumerate(data.items()):
        table.rows[i].cells[0].text = key
        _set_color(table.rows[i].cells[0], key)
        if isinstance(value, dict):
            create_table(table.rows[i].cells[1], value)
        elif isinstance(value, list) and all(isinstance(v, dict) for v in value):
            for v in value:
                create_table(table.rows[i].cells[1], v)
        else:
            _set_color(table.rows[i].cells[1], key)
            table.rows[i].cells[1].text = _round(value)
            if size == COLUMN_SIZE_WITH_DESCRIPTION:
                table.rows[i].cells[2].text = str(description.get(key, "N/A"))


def add_table_caption(doc: Document, caption_text: str) -> None:  # type: ignore
    global TABLE_COUNT  # noqa: PLW0603
    caption = doc.add_paragraph(f"Table {TABLE_COUNT}: {caption_text}", style="Caption")  # type: ignore
    TABLE_COUNT += 1
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.space_after = Pt(0)


def create_tables(
    doc: Document,  # type: ignore
    data: Union[List[Dict[str, Any]], Dict[str, Any]],
    topic: str,
) -> None:
    if isinstance(data, list):
        for d in data:
            add_table_caption(doc, f"Characteristics of {topic} {d['name']}.")
            create_table(doc, d)
            doc.add_paragraph()  # type: ignore

    else:
        add_table_caption(doc, f"Characteristics of {topic} {data['name']}.")
        create_table(doc, data)
        doc.add_paragraph()  # type: ignore


def to_docx(documentation: ModelDocumentation, path: Path) -> None:

    document = Document()
    document.add_page_break()  # type: ignore
    document.add_heading("Spaces", level=2)
    document.add_paragraph(documentation.spaces.introduction)
    create_tables(document, documentation.spaces.table, "Space")
    document.add_paragraph(documentation.spaces.conclusions)

    document.add_page_break()  # type: ignore
    document.add_heading("Constructions", level=2)
    document.add_paragraph(documentation.constructions.introduction)
    create_tables(document, documentation.constructions.table, "Construction")
    document.add_paragraph(documentation.constructions.conclusions)

    document.add_page_break()  # type: ignore
    document.add_heading("Systems", level=2)
    document.add_paragraph(documentation.systems.introduction)
    create_tables(document, documentation.systems.table, "System")
    document.add_paragraph(documentation.systems.conclusions)

    document.save(str(path))
